from telegram import Update
from telegram.ext import ContextTypes

from conversation.session import get_session
from conversation.state import set_state
from conversation.constants import COMPLETED

from documents.complaint_builder import build_complaint

from docx import Document

from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from services.id_generator import generate_complaint_id
from services.case_migration import persist_generated_complaint


# --------------------------------------------------
# DOCX GENERATOR
# --------------------------------------------------

def generate_docx(file_path: str, text: str):
    doc = Document()
    doc.add_heading("Complaint", 0)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(file_path)


# --------------------------------------------------
# PDF GENERATOR
# --------------------------------------------------

def generate_pdf(file_path: str, text: str):
    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    content = []

    for line in text.split("\n"):
        p = Paragraph(line, styles["Normal"])
        content.append(p)
        content.append(Spacer(1, 10))

    doc.build(content)


# --------------------------------------------------
# BUILD TEXT HELPER
# --------------------------------------------------

def build_text(complaint: dict) -> str:
    law = complaint["law"]
    text = []
    text.append(f"Complaint ID: {complaint['complaint_id']}")
    text.append("")
    text.append(f"Date: {complaint['date']}")
    text.append("")
    text.append("Issue:")
    text.append(complaint["issue"])
    text.append("")
    text.append("Legal Ground:")
    text.append(f"{law['law']} - {law['section']}")
    text.append("")
    text.append(law["explanation"])
    return "\n".join(text)


# --------------------------------------------------
# MAIN HANDLER
# --------------------------------------------------

async def handle_generate(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):
    if update.callback_query:
        user_id = update.callback_query.from_user.id
        message = update.callback_query.message
    else:
        user_id = update.effective_user.id
        message = update.message

    session = get_session(user_id)

    if "complaint_id" not in session:
        session["complaint_id"] = generate_complaint_id()

    format_type = session.get("format", "pdf")

    try:
        office = session.get("office", {})
        office_id = office.get("office_id") or office.get("id") or "1"
        identity_mode = session.get("identity_mode", "anonymous")

        if identity_mode == "anonymous":
            user_name = "Anonymous"
            user_address = "Not Provided"
        elif identity_mode == "name_only":
            user_name = session.get("name") or session.get("citizen_name") or "Not Provided"
            user_address = "Not Provided"
        elif identity_mode == "address_only":
            user_name = "Not Provided"
            user_address = session.get("address") or "Not Provided"
        else:
            user_name = session.get("name") or session.get("citizen_name") or "Not Provided"
            user_address = session.get("address") or "Not Provided"

        complaint = build_complaint(
            user_name=user_name,
            user_address=user_address,
            office_id=office_id,
            issue_text=session.get("issue", ""),
        )
        complaint["complaint_id"] = session["complaint_id"]
        complaint_text = build_text(complaint)

        await message.reply_text("Generating document...")

        if format_type == "docx":
            file_path = f"/tmp/complaint_{user_id}.docx"
            filename = "complaint.docx"
            generate_docx(file_path, complaint_text)
        else:
            file_path = f"/tmp/complaint_{user_id}.pdf"
            filename = "complaint.pdf"
            generate_pdf(file_path, complaint_text)

        with open(file_path, "rb") as f:
            await message.reply_document(
                document=f,
                filename=filename,
            )

        await message.reply_text(
            "✅ Document generated for your review/printing. "
            "JanaVani has not submitted it to the government."
        )

        persist_generated_complaint(session)
        set_state(user_id, COMPLETED)

    except Exception as e:
        print("ERROR in handle_generate:", e)
        await message.reply_text(
            "❌ Failed to generate document."
        )
