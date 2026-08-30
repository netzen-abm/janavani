"""Pure PDF/DOCX renderers for reviewed DocumentDraft values."""

from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from src.core.capabilities.document_preparation import DocumentDraft
from src.core.contracts.document_rendering import DocumentFormat, RenderedDocument


class PDFDocumentRenderer:
    format = DocumentFormat.PDF

    def render(self, draft: DocumentDraft, output_path: str) -> RenderedDocument:
        document = SimpleDocTemplate(output_path)
        styles = getSampleStyleSheet()
        story = [Paragraph(escape(draft.document_type.upper()), styles["Title"])]
        story.extend([
            Paragraph(f"<b>To:</b> {escape(draft.to_name)}<br/>{escape(draft.to_address)}", styles["Normal"]),
            Spacer(1, 12),
            Paragraph(f"<b>Subject:</b> {escape(draft.subject)}", styles["Normal"]),
            Spacer(1, 12),
        ])
        if draft.to_email:
            story.append(Paragraph(f"<b>To email:</b> {escape(draft.to_email)}", styles["Normal"]))
        if draft.cc_name or draft.cc_address or draft.cc_email:
            cc = "<b>CC:</b> " + " | ".join(filter(None, [draft.cc_name, draft.cc_address, draft.cc_email]))
            story.append(Paragraph(escape(cc), styles["Normal"]))
            story.append(Spacer(1, 12))
        for paragraph in draft.body.split("\n"):
            story.append(Paragraph(escape(paragraph) if paragraph else "&nbsp;", styles["Normal"]))
        document.build(story)
        return RenderedDocument(draft.document_id, self.format, str(Path(output_path).resolve()))


class DOCXDocumentRenderer:
    format = DocumentFormat.DOCX

    def render(self, draft: DocumentDraft, output_path: str) -> RenderedDocument:
        document = Document()
        document.add_heading(draft.document_type.upper(), level=1)
        document.add_paragraph(f"To: {draft.to_name}\n{draft.to_address}")
        if draft.to_email:
            document.add_paragraph(f"To email: {draft.to_email}")
        if draft.cc_name or draft.cc_address or draft.cc_email:
            document.add_paragraph("CC: " + " | ".join(filter(None, [draft.cc_name, draft.cc_address, draft.cc_email])))
        document.add_paragraph(f"Subject: {draft.subject}")
        for paragraph in draft.body.split("\n"):
            document.add_paragraph(paragraph)
        document.save(output_path)
        return RenderedDocument(draft.document_id, self.format, str(Path(output_path).resolve()))
