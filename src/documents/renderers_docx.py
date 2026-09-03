"""DOCX adapter for the canonical DocumentDraft contract."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from src.documents.document_contract import DocumentDraft


class DocxDocumentRenderer:
    """Render a DocumentDraft to DOCX; never send or submit it."""

    def render(self, draft: DocumentDraft, output_dir: str | Path) -> Path:
        directory = Path(output_dir)
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / f"{draft.document_id}.docx"

        document = Document()
        title = document.add_heading(draft.document_type.upper(), level=1)
        title.alignment = 1
        document.add_paragraph(f"Document ID: {draft.document_id}")
        document.add_paragraph(f"Case ID: {draft.case_id}")
        document.add_paragraph(f"Date: {draft.date}")

        self._add_party(document, "To", draft.to)
        if draft.sender:
            self._add_party(document, "From", draft.sender)

        document.add_paragraph(f"Subject: {draft.subject}").runs[0].bold = True
        document.add_paragraph(draft.body)

        if draft.legal_ground:
            document.add_paragraph("Legal Ground").runs[0].bold = True
            document.add_paragraph(draft.legal_ground)

        if draft.cc:
            document.add_paragraph("CC").runs[0].bold = True
            for party in draft.cc:
                self._add_party(document, party.role or "", party)

        document.save(path)
        return path

    @staticmethod
    def _add_party(document: Document, label: str, party) -> None:
        if label:
            document.add_paragraph(label).runs[0].bold = True
        document.add_paragraph(party.role or party.name)
        if party.role:
            document.add_paragraph(party.name)
        if party.address:
            document.add_paragraph(party.address)
        if party.email:
            document.add_paragraph(f"Email: {party.email}")
