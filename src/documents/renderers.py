"""Canonical document rendering capability.

Renderers are transport- and channel-independent. They consume structured
DocumentPayload values and return in-memory artifacts; they do not access
Telegram, Web, databases, or external AI providers.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from html import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


@dataclass(frozen=True)
class DocumentArtifact:
    """Generated document bytes plus its media type and extension."""

    content: bytes
    media_type: str
    extension: str


class DocumentRenderer:
    """Canonical renderer facade for supported document formats."""

    @staticmethod
    def _lines(payload: dict) -> list[str]:
        title = str(payload.get("title") or "JANAVANI DOCUMENT")
        lines = [title]
        for key, value in payload.items():
            if key == "title" or value in (None, "", []):
                continue
            if isinstance(value, dict):
                lines.append(f"{key.replace('_', ' ').title()}:" )
                lines.extend(f"  {k.replace('_', ' ').title()}: {v}" for k, v in value.items())
            elif isinstance(value, list):
                lines.append(f"{key.replace('_', ' ').title()}: " + "; ".join(map(str, value)))
            else:
                lines.append(f"{key.replace('_', ' ').title()}: {value}")
        return lines

    @classmethod
    def pdf(cls, payload: dict) -> DocumentArtifact:
        buffer = io.BytesIO()
        document = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        story = []
        for index, line in enumerate(cls._lines(payload)):
            if not line.strip():
                story.append(Spacer(1, 10))
            else:
                style = styles["Title"] if index == 0 else styles["BodyText"]
                story.append(Paragraph(escape(line), style))
                story.append(Spacer(1, 6))
        document.build(story)
        return DocumentArtifact(buffer.getvalue(), "application/pdf", ".pdf")

    @classmethod
    def docx(cls, payload: dict) -> DocumentArtifact:
        buffer = io.BytesIO()
        document = Document()
        lines = cls._lines(payload)
        if lines:
            document.add_heading(lines[0], level=1)
            for line in lines[1:]:
                document.add_paragraph(line)
        document.save(buffer)
        return DocumentArtifact(buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx")

    @classmethod
    def render(cls, payload: dict, format_type: str) -> DocumentArtifact:
        normalized = format_type.lower().lstrip(".")
        if normalized == "pdf":
            return cls.pdf(payload)
        if normalized == "docx":
            return cls.docx(payload)
        raise ValueError(f"Unsupported document format: {format_type}")
