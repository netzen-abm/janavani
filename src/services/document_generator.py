import io
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

class MultiFormatDocumentEngine:
    """
    Compiles formal legal documents into print-ready PDF binary streams 
    or editable Microsoft Word (.docx) file formats locally.
    """
    
    @staticmethod
    def generate_pdf_stream(text_content: str) -> io.BytesIO:
        """Renders raw document structures into a professional, printable PDF document stream."""
        pdf_buffer = io.BytesIO()
        
        # Enforce formal structural margins for official government mailings
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Define clean, professional typographic spacing rules
        body_style = ParagraphStyle(
            'FormalBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=16,
            spaceAfter=10
        )
        
        story = []
        # Split text lines safely to maintain vertical layout parity
        lines = text_content.split('\n')
        
        for line in lines:
            if line.strip() == "":
                story.append(Spacer(1, 12))
            else:
                story.append(Paragraph(line, body_style))
                
        doc.build(story)
        pdf_buffer.seek(0)
        return pdf_buffer

    @staticmethod
    def generate_docx_stream(text_content: str) -> io.BytesIO:
        """Renders raw document structures into an editable Word (.docx) file stream."""
        docx_buffer = io.BytesIO()
        doc = Document()
        
        lines = text_content.split('\n')
        for line in lines:
            if line.strip() != "":
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("") # Retain clear whitespace lines
                
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer
